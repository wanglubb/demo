import unittest
import os
from test import parse_doc_to_string

class TestDocParser(unittest.TestCase):
    """测试doc文件解析功能"""

    def setUp(self):
        """测试前的准备工作"""
        # 创建一个测试用的doc文件路径
        self.test_doc_path = "test_document.docx"
        # 如果测试文件不存在，创建一个简单的测试文件
        if not os.path.exists(self.test_doc_path):
            from docx import Document
            doc = Document()
            doc.add_heading("测试文档", 0)
            doc.add_paragraph("这是第一个测试段落。")
            doc.add_paragraph("这是第二个测试段落。")

            # 添加一个表格
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "姓名"
            table.cell(0, 1).text = "年龄"
            table.cell(1, 0).text = "张三"
            table.cell(1, 1).text = "25"

            doc.save(self.test_doc_path)

    def tearDown(self):
        """测试后的清理工作"""
        # 测试完成后删除测试文件
        if os.path.exists(self.test_doc_path):
            os.remove(self.test_doc_path)

    def test_parse_existing_doc(self):
        """测试解析存在的doc文件"""
        result = parse_doc_to_string(self.test_doc_path)

        # 验证结果中包含预期的内容
        self.assertIn("测试文档", result)
        self.assertIn("这是第一个测试段落", result)
        self.assertIn("这是第二个测试段落", result)
        self.assertIn("姓名 | 年龄", result)
        self.assertIn("张三 | 25", result)

    def test_parse_nonexistent_doc(self):
        """测试解析不存在的doc文件"""
        nonexistent_path = "nonexistent_file.docx"
        result = parse_doc_to_string(nonexistent_path)

        # 验证返回了错误信息
        self.assertIn("错误", result)
        self.assertIn(nonexistent_path, result)

    def test_parse_invalid_file(self):
        """测试解析无效的doc文件"""
        # 创建一个非doc文件
        invalid_path = "invalid_file.txt"
        with open(invalid_path, "w") as f:
            f.write("这不是一个doc文件")

        try:
            result = parse_doc_to_string(invalid_path)
            # 验证返回了错误信息
            self.assertIn("错误", result)
        finally:
            # 清理测试文件
            if os.path.exists(invalid_path):
                os.remove(invalid_path)

if __name__ == "__main__":
    unittest.main()
